#!/usr/bin/env python3

# Import necessary libraries
import os
import re
import argparse
import logging
import subprocess
import importlib
import time
from datetime import datetime
from collections import defaultdict
from itertools import combinations
from collections import Counter
import xml.etree.ElementTree as ET

# Mining PubMed
from Bio import Entrez

# Word and Excel
from docx import Document
from docx.shared import Pt, RGBColor, Inches  # Add Pt for font size and color
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Plotting
import seaborn as sns
import networkx as nx
import holoviews as hv
from holoviews import opts, render
from holoviews.plotting.util import dim
import plotly.graph_objects as go
import matplotlib.pyplot as plt

# Data manipulation
import numpy as np
import pandas as pd

# Change log:
# * v1.2.0beta, 2024-11-20: Added collection of pairwise and group collaborations for the found publications. Added new visualizations for collaboration. Added possibility to add more than one organization to search for. Clarified help for arguments. Re-organized the script and added annotations.
# * v1.1.1, 2024-11-19: Improved Word-docx output. Changed logger-output to be less verbose and move things to the --debug. Clarified logger output further. 
# * v1.1.0, 2024-11-18: Fixed an issue where not all the aliases for --names, --departments and --organization were properly queried in conjunction with --organization. Added an option to include ORCID in the author alias list. Fixed issue where the moving average plot might not handle edge years (with fewer than moving_avg_window data points) gracefully.
# * v1.0.10, 2024-11-15: Fixed an issue with consistency of filenaming. Added moving average per author per year to barplot.
# * v1.0.9, 2024-11-15: Fixed an issue with the dimensions of the preprint and publication tables.
# * v1.0.8, 2024-11-15: Fixed issue with aliases for departments. Also edit option to include more departments to search for. Fixed handling authors. Fixed output of preprint citation. 
# * v1.0.7, 2024-11-15: Fixed an issue where the plot for total_publications_preprints_by_author was not displaying the years correctly.
# * v1.0.6, 2024-11-15: Added top 10 journals plot. Fixed issue with JID extraction. Fixed issue with open access extraction. Added more logging. Added --debug flag. 
# * v1.0.5, 2024-11-15: Fixed an issue where the logo was not properly referenced.
# * v1.0.4, 2024-11-15: Added logo to Word document header.
# * v1.0.3, 2024-11-15: Expanded Word-document information.
# * v1.0.2, 2024-11-15: Added retry logic for PubMed API, better logging for aliases, results directory customization, improved input validation, enhanced plotting, and bar annotations.
# * v1.0.1, 2024-11-15: Added alias handling for authors, improved deduplication of YearCount and PubTypeYearCount in Excel, added Authors column.
# * v1.0.0, 2024-11-14: Initial version. Added --year flag for filtering by year range, adjusted tables and figures by author. Stratified tables and figures for each author in DEFAULT_NAMES. Summarized results in Word and Excel files. Added support for "Access Type" in Publications sheet.

# Version and License Information
VERSION_NAME = 'PubMed Miner'
VERSION = '1.2.0beta'
VERSION_DATE = '2024-11-20'
COPYRIGHT_AUTHOR = 'Sander W. van der Laan'
COPYRIGHT = 'Copyright 1979-2024. Sander W. van der Laan | s.w.vanderlaan [at] gmail [dot] com | https://vanderlaanand.science.'
COPYRIGHT_TEXT = '''
Creative Commons Attribution-NonCommercial-NoDerivatives 4.0 International Public License

By exercising the Licensed Rights (defined below), You accept and agree to be bound by the terms and conditions of this Creative Commons Attribution-NonCommercial-NoDerivatives 4.0 International Public License ("Public License"). To the extent this Public License may be interpreted as a contract, You are granted the Licensed Rights in consideration of Your acceptance of these terms and conditions, and the Licensor grants You such rights in consideration of benefits the Licensor receives from making the Licensed Material available under these terms and conditions.

Full license text available at https://creativecommons.org/licenses/by-nc-nd/4.0/legalcode.

This software is provided "as is" without warranties or guarantees of any kind.
'''

# Alias mapping for handling multiple author names and ORCID
ALIAS_MAPPING = {
    "van der Laan SW": [
        "van der Laan SW",
        "van der Laan S", 
        "van der Laan, SW",
        "van der Laan, S",
        "van der Laan, Sander W",
        "van der Laan Sander W",
        "van der Laan, Sander",
        "van der Laan Sander",
        "Sander W van der Laan",
        "Sander van der Laan", 
        "SW van der Laan",
        "S van der Laan",
        "0000-0001-6888-1404",
        # Add other aliases as needed
    ],
    "Pasterkamp G": [
        "Pasterkamp G",
        "Gerard Pasterkamp",
        # Add other aliases as needed
    ],
    "Mokry M": [
        "Mokry M",
        "Michal Mokry",
        # Add other aliases as needed
    ],
    "Schiffelers RM": [
        "Schiffelers RM",
        "Schiffelers R", 
        "Raymond M. Schiffelers", 
        "Raymond Schiffelers", 
        "R. Schiffelers", 
        "R Schiffelers", 
        "Raymond M Schiffelers",
        # Add other aliases as needed
    ],
    "van Solinge W": [
        "van Solinge W",
        "van Solinge WW", 
        "van Solinge W.W.", 
        "Wouter W. van Solinge", 
        "Wouter van Solinge",
        # Add other aliases as needed
    ],
    "Haitjema S": [
        "Haitjema S",
        "Saskia Haitjema",
        "S Haitjema",
        # Add other aliases as needed
    ],
    "den Ruijter HM": [
        "den Ruijter HM",
        "Hester M den Ruijter", 
        "Hester den Ruijter",
        # Add other aliases as needed
    ],
    "Hoefer IE": [
        "Hoefer IE",
        "Imo E Hoefer", 
        "I Hoefer",
        "IE Hoefer",
        "Imo Hoefer",
        "Hofer I", 
        "I Hofer",
        "Imo Hofer", 
        # Add other aliases as needed
    ],
    "Vader P": [
        "Vader P",
        "P Vader", 
        "Pieter Vader",
        "Vader Pieter"
        # Add other aliases as needed
    ],
    # Add other aliases as needed
}

# Departement mapping for handling multiple department names
DEPARTMENT_ALIAS_MAPPING = {
    "Central Diagnostics Laboratory": [
        "Central Diagnostics Laboratory",
        "CDL",
        "CDL Research",
        "Central Diagnostic Laboratory",
        "Central Diagnostic Laboratory Research",
        "Central Diagnostics Laboratory Research",
        "Central Diagnostic Laboratory, Division Laboratories, Pharmacy, and Biomedical genetics",
        "Central Diagnostics Laboratory, Division Laboratories, Pharmacy, and Biomedical genetics"
        "Central Diagnostic Laboratory, Division Laboratory, Pharmacy, and Biomedical genetics",
        "Laboratory of Clinical Chemistry and Hematology, Division Laboratories and Pharmacy",
        "Laboratory of Clinical Chemistry and Hematology, Division Laboratories & Pharmacy",
        "Laboratory of Clinical Chemistry and Hematology",
        "Laboratory Clinical Chemistry and Hematology",
        "LKCH",
        "Laboratory of Clinical Chemistry and Hematology, Division Laboratories and Pharmacy, UMC Utrecht",
        "Laboratorium Klinische Chemie en Hematologie",
        "Laboratorium Klinische Chemie en Hematologie, Divisie Laboratoria en Apotheek",
        "Laboratorium voor Klinische Chemie en Hematologie",
        "Laboratorium voor Klinische Chemie en Hematologie, Divisie Laboratoria en Apotheek",
        "Departments of Clinical Chemistry and Haematology",
        "Department of Clinical Chemistry and Haematology",
        # Add other aliases as needed
    ],
    # Add other departments as needed
    "Laboratory of Experimental Cardiology": [
        "Laboratory of Experimental Cardiology",
        "Laboratory Experimental Cardiology",
        "Lab of Experimental Cardiology",
        "Experimental Cardiology",
        "Experimental Cardiology Lab",
        "Experimental Cardiology Laboratory",
        "Laboratory of Experimental Cardiology, Division Heart and Lungs"
        "Laboratory of Experimental Cardiology, Division of Heart and Lungs",
        "Experimental Cardiology Laboratory Heart Lung Center Utrecht",
        # Add other aliases as needed
    ]
}

# Organization mapping for handling multiple organization names
ORGANIZATION_ALIAS_MAPPING = {
    "University Medical Center Utrecht": [
        "University Medical Center Utrecht",
        "University Medical Center Utrecht, Utrecht University",
        "UMCU",
        "UMC Utrecht",
        "University Medical Centre Utrecht",
        "Universitair Medisch Centrum Utrecht",
        # Add other aliases as needed
    ]
}

# Set some defaults
DEFAULT_NAMES = ["van der Laan SW", 
"Pasterkamp G", 
"Mokry M", 
"Schiffelers RM", 
"van Solinge W", 
"Haitjema S", 
"den Ruijter HM", 
"Hoefer IE",
"Vader P"]
# DEFAULT_DEPARTMENTS = ["Central Diagnostics Laboratory"]
DEFAULT_DEPARTMENTS = ["Central Diagnostics Laboratory", "Laboratory of Experimental Cardiology"]
DEFAULT_ORGANIZATION = ["University Medical Center Utrecht"]

####################################################################################################
#                                   SETUP FUNCTIONS                                                #
####################################################################################################

# Setup Logging
def setup_logger(results_dir, output_base_name, verbose, debug):
    """
    Setup the logger to log to a file and console.
    """
    log_file = os.path.join(results_dir, f"{output_base_name}.log")
    
    os.makedirs(results_dir, exist_ok=True)
    
    logger = logging.getLogger('pubmed_miner')
    # Determine logging level
    if debug:
        level = logging.DEBUG
    elif verbose:
        level = logging.INFO
    else:
        level = logging.WARNING
    logger.setLevel(level)
    
    fh = logging.FileHandler(log_file)
    fh.setLevel(logging.DEBUG)  # File should always capture all messages
    ch = logging.StreamHandler()
    ch.setLevel(level)  # Console output based on verbosity/debug
    
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    fh.setFormatter(formatter)
    ch.setFormatter(formatter)
    
    logger.addHandler(fh)
    logger.addHandler(ch)
    return logger

# Check and install necessary packages
def check_install_package(package_name, logger):
    """
    Check if a package is installed and install it if not.
    """
    try:
        importlib.import_module(package_name)
    except ImportError:
        logger.info(f'{package_name} is not installed. Installing it now...')
        subprocess.check_call(['pip', 'install', package_name])

# Parse year or year range
def parse_year_range(year_range_str):
    """
    Parse year or year range string and return start and end year.
    """
    if '-' in year_range_str:
        start_year, end_year = map(int, year_range_str.split('-'))
    else:
        start_year = end_year = int(year_range_str)
    return start_year, end_year

# Parse command-line arguments
def parse_arguments():
    """
    Parse command-line arguments.
    """
    parser = argparse.ArgumentParser(description=f"""
{VERSION_NAME} v{VERSION} ({VERSION_DATE})
Retrieve PubMed publications for a list of authors.""",
        epilog=f"""
This script retrieves PubMed publications for a list of authors and departments from UMC Utrecht.
It then analyzes the publication data and saves the results to a Word document and an Excel file.

Required arguments:
    -e, --email <email-address>  Email address for PubMed API access.

Optional arguments:
    -n, --names <names>          List of (main) author names to search for. Could also be an ORCID.
                                 Default: {DEFAULT_NAMES} with these aliases: {ALIAS_MAPPING}.
    -dep, --departments <depts>  List of departments to search for. 
                                 Default: {DEFAULT_DEPARTMENTS} with these aliases: {DEPARTMENT_ALIAS_MAPPING}.
    --ignore-departments         Ignore departments in the PubMed query.
    -org, --organization <org>   Organization name for filtering results. 
                                 Default: {DEFAULT_ORGANIZATION} with these aliases {ORGANIZATION_ALIAS_MAPPING}.
    -y, --year <year>            Filter publications by year or year range (e.g., 2024 or 2017-2024).
    -o, --output-file <file>     Output base name for the Word and Excel files. Default: date_CDL_UMCU_Publications.
    -r, --results-dir <dir>      Directory to save results. Default: results.
    --dummy                      Dummy argument for testing. Creates a dummy dataset.
    -v, --verbose                Enable verbose output.
    -d, --debug                  Enable debug output.
    -V, --version                Show program's version number and exit.

Example:
    python pubmed_miner.py --email <email-address> --year 2017-2024 --verbose
+ {VERSION_NAME} v{VERSION}. {COPYRIGHT} +
{COPYRIGHT_TEXT}""",
        formatter_class=argparse.RawTextHelpFormatter)
    parser.add_argument("-e", "--email", required=True, help="Email for PubMed API access. Only used by PubMed to log your query, it is not used to query. Required.")
    parser.add_argument("--dummy", action="store_true", help="Use dummy data for testing and debugging.")
    parser.add_argument("-n", "--names", nargs='+', default=DEFAULT_NAMES, help=f"List of (main) author names or ORCIDs to search for. For example 'van der Laan SW' or '0000-0001-6888-1404'. Defaults: {DEFAULT_NAMES} with aliases {ALIAS_MAPPING}.")
    parser.add_argument("-dep", "--departments", nargs='+', default=DEFAULT_DEPARTMENTS, help=f"List of departments to search for. For example 'Central Diagnostics Laboratory'. Defaults: {DEFAULT_DEPARTMENTS} with aliases {DEPARTMENT_ALIAS_MAPPING}.")
    parser.add_argument("--ignore-departments", action="store_true", help="Ignore departments in the PubMed query.")
    parser.add_argument("-org", "--organization", nargs='+', default=DEFAULT_ORGANIZATION, help=f"Organization name for filtering results. For example 'University Medical Center Utrecht'. Default: {DEFAULT_ORGANIZATION} with aliases {ORGANIZATION_ALIAS_MAPPING}.")
    parser.add_argument("-y", "--year", help="Filter publications by year or year range, for example, 2024 or 2017-2024.")
    parser.add_argument("-o", "--output-file", default="CDL_UMCU_Publications", help="Output base name for the Word and Excel files. Default: date_CDL_UMCU_Publications.")
    parser.add_argument("-r", "--results-dir", default="results", help="Directory to save results. Default: results.")
    parser.add_argument("-v", "--verbose", action="store_true", help="Enable verbose output.")
    parser.add_argument("-d", "--debug", action="store_true", help="Enable debug output. Note: this will produce a lot of output.")
    parser.add_argument('-V', '--version', action='version', version=f'{VERSION_NAME} {VERSION} ({VERSION_DATE})')
    return parser.parse_args()

####################################################################################################
#                                   DATA PROCESSING FUNCTIONS                                      #
####################################################################################################

# Normalize text -- needed for validation_results (could be used elsewhere too)
def normalize_text(text):
    """
    Normalize text for comparison: lowercase, remove punctuation, and normalize spaces.
    """
    import re
    text = text.lower()  # Lowercase
    text = re.sub(r'[^\w\s,\.]', '', text)  # Remove punctuation except commas/periods
    text = re.sub(r'\s+', ' ', text).strip()  # Normalize spaces
    return text.strip()

# Retry logic for API calls
def fetch_with_retry(db, term, retries=3, backoff=2):
    """
    Fetch PubMed data with retry logic to handle rate limits or network issues.
    """
    for attempt in range(retries):
        try:
            handle = Entrez.esearch(db=db, term=term, retmax=100)
            return Entrez.read(handle)
        except Exception as e:
            if attempt < retries - 1:
                time.sleep(backoff ** attempt)
                continue
            raise e

# Retry logic for API calls -- needed for the process_collaborations function
def fetch_pubmed_metadata(pubmed_id, email, logger):
    """
    Fetch and return metadata for a PubMed article, specifically extracting authors.

    Args:
        pubmed_id (str): PubMed ID of the article.
        email (str): Email address to use for PubMed API.
        logger (Logger): Logger object for logging.

    Returns:
        list: List of authors for the given PubMed ID.
    """
    Entrez.email = email

    try:
        logger.debug(f"Fetching metadata for PubMed ID {pubmed_id}.")
        # Fetch metadata in MEDLINE format
        with Entrez.efetch(db="pubmed", id=pubmed_id, rettype="medline", retmode="text") as handle:
            record = handle.read()

        # Extract authors from the MEDLINE format
        logger.debug(f"Parsing authors for PubMed ID {pubmed_id}.")
        authors = []
        for line in record.split("\n"):
            if line.startswith("AU  -"):  # Author line in MEDLINE format
                author_name = line.replace("AU  - ", "").strip()
                authors.append(author_name)

        if not authors:
            logger.warning(f"No authors found for PubMed ID {pubmed_id}.")
        else:
            logger.debug(f"Authors found for PubMed ID {pubmed_id}: {authors}")

        return authors
    except Exception as e:
        logger.error(f"Error fetching metadata for PubMed ID {pubmed_id}: {e}")
        return []

# Extract affiliations -- needed in validation_results
def extract_affiliations(record):
    """
    Extract all affiliations (AD fields) from a MEDLINE record, accounting for multi-line entries.
    """
    affiliations = []
    lines = record.splitlines()

    # Flag to track if we're in an "AD" block
    in_ad_block = False
    current_affiliation = ""

    for line in lines:
        if line.startswith("AD  - "):  # Start of a new affiliation
            if current_affiliation:  # If there's an ongoing affiliation, save it
                affiliations.append(current_affiliation.strip())
            current_affiliation = line[6:]  # Strip "AD  - " prefix
            in_ad_block = True
        elif in_ad_block and line.startswith("      "):  # Continuation line for "AD"
            current_affiliation += " " + line.strip()  # Append the continuation
        else:
            if in_ad_block:  # End of "AD" block
                affiliations.append(current_affiliation.strip())
                current_affiliation = ""
                in_ad_block = False

    # Catch any remaining affiliation
    if current_affiliation:
        affiliations.append(current_affiliation.strip())

    return " ".join(affiliations)  # Combine into a single string

# Get canonical author name
def get_canonical_author(author, logger=None):
    """
    Get the canonical author name or ORCID from the ALIAS_MAPPING.
    First it checks for ORCID format.
    Next, match given author to any known aliases and return the canonical name
    Last, return the input author if no match is found among the aliases.
    """

    # Check for ORCID format
    # If the input author is in the format of an ORCID (e.g., 0000-0001-6888-1404), 
    # it returns the input directly, assuming it's already canonical.
    if re.match(r"^\d{4}-\d{4}-\d{4}-\d{4}$", author):  # Check for ORCID format
        if logger:
            logger.info(f"Processing ORCID: {author}.")
        return author  # Return ORCID directly

    # Search for the author in the aliases
    # For other inputs, it loops through all canonical authors and their aliases in ALIAS_MAPPING:
    # - If the input matches any alias, it returns the corresponding canonical name.
    # - If the input matches a canonical name directly, it also returns the canonical name.
    # - If no match is found, it returns the input author as is.
    
    for canonical, aliases in ALIAS_MAPPING.items():
        if logger:
            logger.debug(f"Checking aliases for canonical '{canonical}': {aliases}")
        if author in aliases:
            if logger:
                logger.debug(f"> Matched author alias given, '{author}', to the canonical author: '{canonical}'.")
            return canonical
        if author == canonical: # Exact match with a canonical name
            if logger:
                logger.debug(f"> Alias given is the canonical author: '{author}'.")
            return canonical
    
    # Fallback: return the input author if no match is found among the known aliases
    if logger:
        logger.warning(f"> No match among aliases found for author '{author}'. Returning as is. Note: usually this means the author is not in the ALIAS_MAPPING; depending on the use-case, this is normally not an issue to worry about (you can add more aliases if you want).")
    return author

# Fetch publication detailss
def fetch_publication_details(pubmed_ids, logger, main_author, start_year=None, end_year=None):
    """
    Fetch detailed information for each PubMed ID, filter by year, and identify preprints.
    """
    canonical_author = get_canonical_author(main_author, logger)
    publications = []
    preprints = []

    for pub_id in pubmed_ids:
        record = None

        # Retry logic for Entrez.efetch
        for attempt in range(3):
            try:
                handle = Entrez.efetch(db="pubmed", id=pub_id, rettype="medline", retmode="text")
                record = handle.read()
                handle.close()
                break  # Exit loop on success
            except Exception as e:
                if attempt < 2:  # Retry for first two attempts
                    time.sleep(2 ** attempt)  # Exponential backoff
                    logger.warning(f"Retrying PubMed fetch for ID {pub_id} (attempt {attempt + 2})...")
                else:  # Final failure
                    logger.error(f"Failed to fetch PubMed details for ID {pub_id}: {e}")
                    continue

        if not record:
            continue  # Skip this ID if all retries failed

        # Extract publication details
        authors = re.findall(r"AU  - (.+)", record) or []  # Short author list
        full_authors = re.findall(r"FAU - (.+)", record) or []  # Full author list
        author_ids = re.findall(r"AUID- (.+)", record) or []  # Author identifiers (e.g., ORCID)

        # Replace aliases with canonical names for `AU` and `FAU` fields
        authors = [get_canonical_author(author) for author in authors]
        full_authors = [get_canonical_author(author) for author in full_authors]

        # Combine `AU`, `FAU`, and `AUID` into a single list for alias matching
        all_author_data = set(authors + full_authors + author_ids)

        # Check if the canonical author or any alias matches
        author_match = any(alias in all_author_data for alias in all_author_data)

        title = re.search(r"TI  - (.+)", record).group(1) if re.search(r"TI  - (.+)", record) else "No title found"
        journal_abbr = re.search(r"TA  - (.+)", record).group(1) if re.search(r"TA  - (.+)", record) else "No journal abbreviation found"
        # Extract Journal ID with debugging print
        jid_match = re.search(r"JID\s*-\s*(\d+)", record)
        if jid_match:
            journal_id = jid_match.group(1)
            # debug
            # logger.debug(f"Extracted Journal ID: {journal_id}")
        else:
            journal_id = "No journal ID"
            # debug
            # logger.debug(f"No Journal ID found in record.")
        pub_date = re.search(r"DP  - (.+)", record).group(1)[:4] if re.search(r"DP  - (.+)", record) else "Unknown Year"
        doi_match = re.search(r"AID - (10\..+?)(?: \[doi\])", record)
        doi_link = f"https://doi.org/{doi_match.group(1)}" if doi_match else "No DOI found"
        pub_type = re.findall(r"PT  - (.+)", record)
        source = re.search(r"SO  - (.+)", record).group(1) if re.search(r"SO  - (.+)", record) else "No source found"

        # Extract UOF (Update of) field
        uof_match = re.search(r"UOF\s*-\s*(.+)", record)
        uof = uof_match.group(1) if uof_match else "No UOF"
        # Remove "PMID:" from the uof field, if it exists
        uof = uof.replace("PMID:", "").strip()

        # Determine publication type based on PT and SO
        publication_type = "Other"
        if "Review" in pub_type or "Review" in source:
            publication_type = "Review"
        elif "Book" in pub_type or "Book" in source:
            publication_type = "Book"
        elif "Journal Article" in pub_type or "Journal Article" in source:
            publication_type = "Journal Article"

        # Determine access type based on the presence of a PMC ID
        pmc_match = re.search(r"PMC\s+-\s+PMC\d+", record)
        if pmc_match:
            access_type = "open access"
            # debug
            # logger.debug(f"Detected PMC ID: {pmc_match.group()}")
        else:
            access_type = "closed access"
            # debug
            # logger.debug(f"No PMC ID detected; setting access type to 'closed access'.")

        # Skip errata or corrections
        if "ERRATUM" in title.upper() or "AUTHOR CORRECTION" in title.upper():
            logger.info(f"Skipping 'erratum' or 'author correction' for [{title}]")
            continue

        # Filter by year if specified
        if start_year and end_year:
            if not (start_year <= int(pub_date) <= end_year):
                logger.info(f"Skipping [{title}] as it falls outside the year range.")
                continue
        
        # Check if the UOF references preprint servers and extract DOI
        uof_doi = None
        if any(preprint_source in uof for preprint_source in ["medRxiv", "bioRxiv", "arXiv"]):
            logger.info(f"PMID {pub_id} references a preprint ({uof}). Adding to preprints.")
            uof_parts = uof.split()
            uof_doi = None
            for part in uof_parts:
                if part.startswith("10."):
                    uof_doi = f"https://doi.org/{part}"
                    break
            uof_citation = " ".join(uof_parts[:uof_parts.index(part)]) if uof_doi else uof

            # Remove "doi:" from the uof_citation if present
            if "doi:" in uof_citation.lower():
                uof_citation = re.sub(r"\bdoi:\b", "", uof_citation, flags=re.IGNORECASE).strip()

            preprints.append(
                (
                    pub_id,
                    canonical_author if canonical_author in authors else f"{canonical_author} et al.",
                    pub_date,
                    journal_abbr,
                    journal_id,
                    title,
                    uof_doi,
                    uof_citation,
                    publication_type,
                )
            )
        # Separate preprints
        if "Preprint" in pub_type:
            preprints.append((pub_id, canonical_author if canonical_author in authors else f"{canonical_author} et al.", pub_date, journal_abbr, journal_id, title, doi_link, source, publication_type))
        else:
            publications.append((pub_id, canonical_author if canonical_author in authors else f"{canonical_author} et al.", pub_date, journal_abbr, journal_id, title, doi_link, source, publication_type, access_type))
        # debug
        logger.debug(f"Found MEDLINE record for PubMed ID {pub_id} matching criteria (within year range).")
        logger.debug(f"Authors: {authors}")
        logger.debug(f"Title: {title}")
        logger.debug(f"Journal: {journal_abbr}")
        logger.debug(f"Journal ID: {journal_id}")
        logger.debug(f"Publication Date: {pub_date}")
        logger.debug(f"DOI: {doi_link}")
        logger.debug(f"Publication Type: {publication_type}")
        logger.debug(f"Citation: {source}")
        logger.debug(f"Access Type: {access_type}")
        logger.debug(f"UOF: {uof}")
        # debug - this produces a lot of output
        # logger.debug(f"This was the full record:\n{record}")

    return publications, preprints

# Analyze publication data
def analyze_publications(publications_data, main_author, logger):
    """
    Analyze the publication data and return counts for authors, years, and journals.
    Also return counts of publications by author per year and by year per journal.
    """
    canonical_author = get_canonical_author(main_author, logger)
    author_count = defaultdict(int)
    year_count = defaultdict(int)
    author_year_count = defaultdict(lambda: defaultdict(int))
    year_journal_count = defaultdict(lambda: defaultdict(int))
    pub_type_count = defaultdict(lambda: defaultdict(int))

    for pub in publications_data:
        (
            pub_id,
            author,
            year,
            journal_abbr,
            journal_id,
            title,
            doi_link,
            citation,
            pub_type,
            access_type,
        ) = pub  # Adjusted unpacking to match the expanded structure
        
        if canonical_author in author:
            author_count[canonical_author] += 1
            author_year_count[canonical_author][year] += 1
            year_count[year] += 1
            year_journal_count[year][journal_abbr] += 1
            pub_type_count[pub_type][year] += 1
    
    return author_count, year_count, author_year_count, year_journal_count, pub_type_count

# Process collaborations between authors based on publications
def process_collaborations(author_data, logger, alias_mapping, results_dir, output_base_name, email, max_group_size=5):
    """
    Process collaborations by fetching publication details from Entrez and analyzing author collaborations.

    Args:
        author_data (dict): Contains author publication data with PubMed IDs.
        logger: Logger instance.
        alias_mapping (dict): Mapping of canonical authors to their aliases.
        results_dir (str): Directory to save results.
        output_base_name (str): Base name for output files.
        email (str): Email address for Entrez queries.
        max_group_size (int): Maximum group size to consider for collaborations.

    Returns:
        tuple:
            collaboration_data (list): Pairwise collaborations with details.
            group_collaboration_data (Counter): Group collaborations with counts.
            collaboration_matrix (np.array): Matrix showing pairwise collaboration counts.
            authors (list): List of unique canonical authors.
    """
    from Bio import Entrez
    Entrez.email = email  # Set email for Entrez queries

    alias_to_canonical = {
        alias: canonical
        for canonical, aliases in alias_mapping.items()
        for alias in aliases
    }

    # Initialize structures
    collaboration_counts = defaultdict(lambda: defaultdict(int))
    group_collaboration_data = Counter()
    collaboration_details = []
    group_details = []
    canonical_authors = set(alias_to_canonical.values())

    logger.info(f"> Processing collaborations for {len(author_data)} authors.")

    for canonical_author, (publications, preprints, _, _, _, _, _) in author_data.items():
        all_publications = publications + preprints
        logger.info(f"> Processing {len(all_publications)} records for author '{canonical_author}'.")

        for pub in all_publications:
            pub_id, _, year, journal, jid, title, doi, citation, *_ = pub

            # Fetch record details using fetch_collabs_with_retry
            try:
                logger.debug(f"Fetching record details for PMID {pub_id}.")
                authors = fetch_pubmed_metadata(pub_id, email, logger)  # Custom function to fetch authors
                if not authors:
                    logger.warning(f"No authors found for PubMed ID {pub_id}. Skipping.")
                    continue

                # Map authors to canonical names
                publication_authors = {
                    alias_to_canonical.get(author.strip(), None)
                    for author in authors
                }
                publication_authors.discard(None)  # Remove None values

                if len(publication_authors) < 2:
                    logger.debug(f"Skipping publication with less than 2 canonical authors: PubMed ID {pub_id}")
                    continue

                # Pairwise collaborations
                for author1, author2 in combinations(sorted(publication_authors), 2):
                    collaboration_counts[author1][author2] += 1
                    collaboration_counts[author2][author1] += 1
                    collaboration_details.append(
                        (author1, author2, pub_id, year, journal, jid, title, doi, citation)
                    )

                # Group collaborations
                for group_size in range(3, min(len(publication_authors), max_group_size) + 1):
                    for group in combinations(sorted(publication_authors), group_size):
                        group_collaboration_data[group] += 1
                        group_details.append(
                            (group, pub_id, year, journal, jid, title, doi, citation)
                        )
            except Exception as e:
                logger.error(f"Error fetching PubMed record for {pub_id}: {e}")
                continue

    # Create pairwise collaborations matrix
    logger.info("> Creating collaboration matrix.")
    collaboration_data = [
        (author1, author2, count)
        for author1, collabs in collaboration_counts.items()
        for author2, count in collabs.items()
    ]

    authors = sorted(canonical_authors)
    author_index = {author: idx for idx, author in enumerate(authors)}
    collaboration_matrix = np.zeros((len(authors), len(authors)), dtype=int)

    for author1, author2, count in collaboration_data:
        idx1, idx2 = author_index[author1], author_index[author2]
        collaboration_matrix[idx1, idx2] = count
        collaboration_matrix[idx2, idx1] = count

    # Log the collaboration matrix
    logger.debug("Collaboration matrix created:")
    logger.debug(pd.DataFrame(collaboration_matrix, index=authors, columns=authors))

    # Create DataFrames
    logger.info("Creating dataset with pairwise collaboration details.")
    pairwise_df = pd.DataFrame(
        collaboration_details,
        columns=["Canonical Author 1", "Canonical Author 2", "PMID", "Year",
                 "Journal", "JID", "Title", "DOI Link", "Citation"]
    )
    logger.info("Creating dataset with group collaboration details.")
    group_df = pd.DataFrame(
        group_details,
        columns=["Group", "PMID", "Year", "Journal", "JID", "Title", "DOI Link", "Citation"]
    )
    group_df["Group"] = group_df["Group"].apply(lambda x: ", ".join(x))

    # Save to a single Excel file with two sheets but only if there are collaborations
    if not collaboration_data:
        logger.warning("No collaboration data found. Skipping Excel file creation.")
        return collaboration_data, group_collaboration_data, collaboration_matrix, authors
    else:
        logger.info("Saving collaboration details to Excel file.")
        excel_path = os.path.join(results_dir, f"{output_base_name}_collaborations.xlsx")
        with pd.ExcelWriter(excel_path, engine='xlsxwriter') as writer:
            pairwise_df.to_excel(writer, sheet_name="Pairwise_Collab", index=False)
            group_df.to_excel(writer, sheet_name="Group_Collab", index=False)

        logger.info(f"Collaboration details saved to {excel_path}.")

    return collaboration_data, group_collaboration_data, collaboration_matrix, authors

# Generate dummy data for testing
def generate_dummy_data():
    dummy_authors = ["Author A", "Author B", "Author C", "Author D", "Author E"]
    dummy_journals = ["Journal 1", "Journal 2", "Journal 3"]
    dummy_years = range(2015, 2023)
    dummy_data = defaultdict(lambda: ([], [], {}, {}, {}, {}, {}))  # Ensure all fields match expectations

    for i in range(500):  # Create 500 dummy publications
        pub_id = f"DUMMY{i+1}"
        year = np.random.choice(dummy_years)
        journal = np.random.choice(dummy_journals)
        title = f"Dummy Title {i+1}"
        authors = np.random.choice(dummy_authors, size=np.random.randint(2, 4), replace=False).tolist()
        doi = f"https://doi.org/{i+1}"
        citation = f"Dummy Citation {i+1}"
        access_type = np.random.choice(["open access", "closed access"])
        pub_type = np.random.choice(["Journal Article", "Review", "Other"])
        
        publication = (pub_id, ", ".join(authors), year, journal, f"JID{i % 10}", title, doi, citation, pub_type, access_type)

        # Assign publication to a random canonical author
        canonical_author = np.random.choice(dummy_authors)
        dummy_data[canonical_author][0].append(publication)  # Add to publications

        # Add preprints with a smaller probability
        if np.random.rand() < 0.2:  # 20% chance
            preprint = (pub_id, ", ".join(authors), year, journal, f"JID{i % 10}", title, doi, citation, pub_type)
            dummy_data[canonical_author][1].append(preprint)  # Add to preprints

    # Generate counts for each canonical author
    for canonical_author in dummy_authors:
        publications, preprints, _, _, _, _, _ = dummy_data[canonical_author]
        author_count = defaultdict(int)
        year_count = defaultdict(int)
        author_year_count = defaultdict(lambda: defaultdict(int))
        year_journal_count = defaultdict(lambda: defaultdict(int))
        pub_type_count = defaultdict(lambda: defaultdict(int))

        for pub in publications:
            _, author_list, year, journal, _, _, _, _, pub_type, _ = pub
            author_list = author_list.split(", ")
            for author in author_list:
                author_count[author] += 1
                author_year_count[author][year] += 1
            year_count[year] += 1
            year_journal_count[year][journal] += 1
            pub_type_count[pub_type][year] += 1

        dummy_data[canonical_author] = (
            publications,
            preprints,
            author_count,
            year_count,
            author_year_count,
            year_journal_count,
            pub_type_count,
        )

    return dummy_data

####################################################################################################
#                                   WRITE RESULTS TO EXCEL                                         #
####################################################################################################

# Write results to Excel
def write_to_excel(author_data, output_base_name, results_dir, logger):
    """
    Write the combined results for all authors into six sheets of an Excel file.
    """
    logger.info("Writing results to Excel file.")
    # output_file = f"{datetime.now().strftime('%Y%m%d')}_{output_file}"  # Add date to the filename
    # writer = pd.ExcelWriter(os.path.join(results_dir, f"{output_file}.xlsx"), engine='xlsxwriter')
    output_path = os.path.join(results_dir, f"{output_base_name}.xlsx")
    writer = pd.ExcelWriter(output_path, engine='xlsxwriter')

    # Combine all publications into a single DataFrame
    logger.info("> Combining all publications into a single DataFrame.")
    publications_data = []
    for main_author, (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count) in author_data.items():
        canonical_author = get_canonical_author(main_author, logger)
        for pub in publications:
            publications_data.append(list(pub) + [canonical_author])
    logger.debug(f"Publications data: {publications_data[:3]}")  # Log first 3 publications
    publications_df = pd.DataFrame(
        publications_data,
        columns=["PubMed ID", "Author", "Year", "Journal", "JID", "Title", "DOI Link", "Citation", "Publication Type", "Access Type", "Main Author"]
    )

    # Deduplicate based on PubMed ID and combine authors for duplicates in publications
    publications_df = (
        publications_df.groupby("PubMed ID")
        .agg({
            "Author": lambda x: ", ".join(sorted(set(x))),  # Combine full unique author names
            "Year": "first",
            "Journal": "first",
            "JID": "first",
            "Title": "first",
            "DOI Link": "first",
            "Citation": "first",
            "Publication Type": "first",
            "Access Type": "first",
        })
        .reset_index()
    )

    publications_df["Year"] = pd.to_numeric(publications_df["Year"], errors='coerce')
    publications_df.to_excel(writer, sheet_name="Publications", index=False)

    # Combine all preprints into a single DataFrame with deduplication
    logger.info("> Combining all preprints into a single DataFrame with deduplication.")
    preprints_data = []
    for main_author, (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count) in author_data.items():
        canonical_author = get_canonical_author(main_author, logger)
        for preprint in preprints:
            preprints_data.append(list(preprint) + [canonical_author])
    logger.debug(f"Preprints data: {preprints_data[:3]}")  # Log first 3 preprints
    preprints_df = pd.DataFrame(
        preprints_data,
        columns=["PubMed ID", "Author", "Year", "Journal", "JID", "Title", "DOI Link", "Citation", "Publication Type", "Main Author"]
    )

    # Deduplicate based on PubMed ID and combine authors for duplicates in preprints
    preprints_df = (
        preprints_df.groupby("PubMed ID")
        .agg({
            "Author": lambda x: ", ".join(sorted(set(x))),  # Combine full unique author names
            "Year": "first",
            "Journal": "first",
            "JID": "first",
            "Title": "first",
            "DOI Link": "first",
            "Citation": "first",
            "Publication Type": "first",
        })
        .reset_index()
    )

    preprints_df["Year"] = pd.to_numeric(preprints_df["Year"], errors='coerce')
    preprints_df.to_excel(writer, sheet_name="Preprints", index=False)

    # Combine author counts
    logger.info("> Combining author counts.")
    author_counts = []
    for main_author, (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count) in author_data.items():
        canonical_author = get_canonical_author(main_author, logger)
        for author, count in author_count.items():
            author_counts.append([author, count, canonical_author])
    author_counts_df = pd.DataFrame(author_counts, columns=["Author", "Number of Publications", "Main Author"])
    author_counts_df.drop(columns=["Main Author"], inplace=True)
    author_counts_df.to_excel(writer, sheet_name="AuthorCount", index=False)

    # Combine and deduplicate year counts
    logger.info("> Combining and deduplicating year counts.")
    year_counts = []
    for main_author, (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count) in author_data.items():
        canonical_author = get_canonical_author(main_author, logger)
        for year, count in year_count.items():
            year_counts.append([year, count, canonical_author])
    year_counts_df = pd.DataFrame(year_counts, columns=["Year", "Number of Publications", "Authors"])

    year_counts_df = (
        year_counts_df.groupby("Year")
        .agg({
            "Number of Publications": "sum",
            "Authors": lambda x: ", ".join(sorted(set(x)))
        })
        .reset_index()
    )

    year_counts_df["Year"] = pd.to_numeric(year_counts_df["Year"], errors='coerce')
    year_counts_df.to_excel(writer, sheet_name="YearCount", index=False)

    # Combine and deduplicate publication type by year counts
    logger.info("> Combining and deduplicating publication type by year counts.")
    pub_type_year_counts = []
    for main_author, (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count) in author_data.items():
        canonical_author = get_canonical_author(main_author, logger)
        for pub_type, years in pub_type_count.items():
            for year, count in years.items():
                pub_type_year_counts.append([pub_type, year, count, canonical_author])
    pub_type_year_df = pd.DataFrame(pub_type_year_counts, columns=["Publication Type", "Year", "Number of Publications", "Authors"])

    pub_type_year_df = (
        pub_type_year_df.groupby(["Publication Type", "Year"])
        .agg({
            "Number of Publications": "sum",
            "Authors": lambda x: ", ".join(sorted(set(x)))
        })
        .reset_index()
    )

    pub_type_year_df["Year"] = pd.to_numeric(pub_type_year_df["Year"], errors='coerce')
    pub_type_year_df.to_excel(writer, sheet_name="PubTypeYearCount", index=False)

    # Save the Excel file
    logger.info(f"Excel file saved to [{output_path})].\n")
    writer.close()
    
####################################################################################################
#                                   WRITE RESULTS TO WORD                                          #
####################################################################################################

# Word document table creation
def add_table_to_doc(doc, data, headers, title):
    """
    Add a table to a Word document.
    Safeguard against empty or malformed rows.
    data: List of rows, each row is a list of items.
    headers: List of column headers.
    title: Title for the table.
    """

    # Access the document's styles
    styles = doc.styles

    heading3_style = styles['Heading 3']
    heading3_style.font.name = 'Helvetica'
    heading3_style.font.size = Pt(11)
    heading3_style.font.color.rgb = RGBColor(89,90,92)  # Custom color (RGB) 89,90,92 Grey - Uithof Color

    # Add the title for the table
    doc.add_paragraph(title, style='Heading 3')

    # Create the table with headers
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header

    # Add data rows
    for row in data:
        # Skip empty or malformed rows
        if not row or len(row) != len(headers):
            continue
        row_cells = table.add_row().cells
        for idx, item in enumerate(row):
            row_cells[idx].text = str(item)
    
    # Add a blank paragraph for spacing after the table
    doc.add_paragraph()

# Write results to Word
def write_to_word(author_data, output_base_name, results_dir, logger, args):
    """
    Write the combined results for all authors to a Word document.
    """
    logger.info("Writing results to Word document.")
    # output_file = f"{datetime.now().strftime('%Y%m%d')}_{output_file}"  # Add date to the filename
    query_date = datetime.now().strftime('%Y-%m-%d')
    query_year = datetime.now().year
    query_quarter = (datetime.now().month - 1) // 3 + 1
    document = Document()

    # Add a header for the logo
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]

    # Add the logo to the header
    logo_path = os.path.join("images/FullLogo_Transparent.png")  # Ensure path is correct
    try:
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.5))  # Adjust size as needed
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    except Exception as e:
        logger.error(f"Could not add logo to Word document: {e}")

    # Access the document's styles
    styles = document.styles

    # Set the font for the Normal style (base for most text)
    normal_style = styles['Normal']
    normal_style.font.name = 'Helvetica'
    normal_style.font.size = Pt(11)  # Optional: set default font size

    # Set the font for Title, Heading, and other styles
    title_style = styles['Title']
    title_style.font.name = 'Helvetica'
    title_style.font.size = Pt(24)
    title_style.font.color.rgb = RGBColor(18, 144, 217)  # Custom color (RGB) 18,144,217 Azurblue - Uithof Color

    subtitle_style = styles['Subtitle']
    subtitle_style.font.name = 'Helvetica'
    subtitle_style.font.size = Pt(18)
    subtitle_style.font.color.rgb = RGBColor(19, 150, 216)  # Custom color (RGB) 19,150,216 Light Azurblue - Uithof Color

    heading1_style = styles['Heading 1']
    heading1_style.font.name = 'Helvetica'
    heading1_style.font.size = Pt(14)
    heading1_style.font.color.rgb = RGBColor(47,139,201)  # Custom color (RGB) 47,139,201 Skyblue - Uithof Color

    heading2_style = styles['Heading 2']
    heading2_style.font.name = 'Helvetica'
    heading2_style.font.size = Pt(12)
    heading2_style.font.color.rgb = RGBColor(21,166,193)  # Custom color (RGB) 21,166,193 Greenblue - Uithof Color

    heading3_style = styles['Heading 3']
    heading3_style.font.name = 'Helvetica'
    heading3_style.font.size = Pt(11)
    heading3_style.font.color.rgb = RGBColor(89,90,92)  # Custom color (RGB) 89,90,92 Grey - Uithof Color

    bullet_style = styles['List Bullet']
    bullet_style.font.name = 'Helvetica'
    bullet_style.font.size = Pt(11)

    bullet2_style = styles['List Bullet 2']
    bullet2_style.font.name = 'Helvetica'
    bullet2_style.font.size = Pt(11)

    # Add document meta data
    document.core_properties.title = f"Publications for {query_year}-Q{query_quarter}"
    document.core_properties.author = f"{COPYRIGHT_AUTHOR}"
    document.core_properties.comments = f"{VERSION_NAME} v{VERSION} ({VERSION_DATE})"

    # Add main content to the document
    # Add title and subtitle
    document.add_paragraph(f"Publications for {query_year}-Q{query_quarter}", style='Title')
    document.add_paragraph(
        f"Summarizing publications from {DEFAULT_DEPARTMENTS} at {DEFAULT_ORGANIZATION}.", 
        style='Subtitle'
    )

    # Add a heading
    document.add_paragraph("Introduction", style='Heading 1')

    # Add the introductory paragraph
    document.add_paragraph(
        (
            "This report was created using PubMed Miner, a tool to retrieve and analyze publication data from PubMed. "
            "In short, the script mines PubMed through the Entrez API, retrieves the data, and analyzes it to provide insights "
            "into the publications of the CDL.\n"
            "It uses the given author names, departments, and organization to filter the results. In this version, a few "
            "default authors, departments, organizations, and their aliases are provided. If you think this list is incomplete, please "
            "contact us and we will add your name. Nowadays, PubMed also stores the OrcID for a given author, so you can also use this to identify "
            "authors.\n"
            "For our purposes, we are interested in the output for a given year, which can be indicated by the flag `--year`, "
            "for example, `--year 2024`. Alternatively, you could query within a certain timeframe, for example `--year 2022-2024`. "
            "You can also omit this flag, and it will mine PubMed for any publication across all years. Note that the latter "
            "would significantly slow down the process, especially when multiple authors are given."
            "Note that you have to use the `--email` flag to provide your email address for the Entrez API. This is required "
            "to access PubMed data. If you do not provide this, the script will not run. This email is only used by PubMed to "
            "log the queries, it is not used for the actual query.\n"
            "Lastly, the script will save the results in a Word document and an Excel file. The Word document will contain the "
            "main publications and preprints for each author, as well as summary tables for the number of publications per author, "
            "per year, per year per journal, and per publication type. The Excel file will contain the same data but in separate "
            "sheets for each type of data. The Excel file will also contain the raw data for the publications and preprints. The "
            "plots are also saved in the results directory. In addition, the script will log the results and any errors in a log "
            "file, which will be saved in the results directory."

        ), 
        style='Normal'
    )

    # Add a main heading
    document.add_paragraph("The following settings are used:", style='Normal')

    # Add the first-level bullets
    document.add_paragraph(f"Query date: {query_date}.", style="List Bullet")
    document.add_paragraph(f"Authors: {', '.join(author_data.keys())}.", style="List Bullet")

    # Add a second-level bullet for "Author aliases used" only if the user provided author are in the default list
    if set(author_data.keys()) == set(DEFAULT_NAMES):
        document.add_paragraph(f"Author aliases used: {', '.join(ALIAS_MAPPING.keys())}.", style="List Bullet 2")
    else:
        document.add_paragraph(f"No author aliases known.", style="List Bullet 2")

    # Continue with first-level bullets
    document.add_paragraph(f"Year range: {args.year}." if args.year else "No year filter used.", style="List Bullet")
    document.add_paragraph(f"Department(s): {args.departments}.", style="List Bullet")

    # Add a second-level bullet for "Department aliases used" only if the user provided departments are in the default list
    if args.departments == DEFAULT_DEPARTMENTS:
        document.add_paragraph(f"Department aliases used: {', '.join(DEPARTMENT_ALIAS_MAPPING.keys())}.", style="List Bullet 2")
    else:
        document.add_paragraph(f"No department aliases known.", style="List Bullet 2")

    # Continue with first-level bullets
    document.add_paragraph(f"Organization(s): {args.organization}.", style="List Bullet")

    # Add a second-level bullet for "Organization aliases used" only if the user provided organization is the default
    if args.organization == DEFAULT_ORGANIZATION:
        document.add_paragraph(f"Organization aliases used: {', '.join(ORGANIZATION_ALIAS_MAPPING.keys())}.", style="List Bullet 2")
    else:
        document.add_paragraph(f"No organization aliases known.", style="List Bullet 2")

    # Add Document Summary
    document.add_paragraph("Summary", style='Heading 1')
    document.add_paragraph(
        f"Total authors processed: {len(author_data)}\n"
        f"Total publications: {sum(len(data[0]) for data in author_data.values())}\n"
        f"Total preprints: {sum(len(data[1]) for data in author_data.values())}",
        style='Normal'
    )
    document.add_paragraph()
    document.add_paragraph(f"Results saved on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.", style='Normal')
    document.add_paragraph(f"Log file saved to {os.path.join(results_dir, f'{output_base_name}.log')}.", style='Normal')
    document.add_paragraph()
    document.add_paragraph(f"{VERSION_NAME} v{VERSION} ({VERSION_DATE}).", style='Normal')
    document.add_paragraph(f"{COPYRIGHT}", style='Normal')
    document.add_paragraph()
    document.add_paragraph(f"GitHub repository: https://github.com/swvanderlaan/PubMed_Miner. \nAny issues or requests? Create one here: https://github.com/swvanderlaan/PubMed_Miner/issues.", style='Normal')

    # Add a heading for the results
    document.add_paragraph("Results", style='Heading 1')
    # Add results for each canonical author
    logger.info(f"> Adding results for {len(author_data)} author(s).")
    for canonical_author, (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count) in author_data.items():
        # Add a heading for the author
        document.add_paragraph(f"Author: {canonical_author}", style='Heading 2')

        # Main publications and preprints for this author
        if publications:
            logger.info(f"> Adding {len(publications)} publication(s) for {canonical_author}.")
            
            # Debugging: Check structure of publications
            logger.debug(f"Publications for {canonical_author}: {publications[:3]}")  # Log first 3 publications
            
            # Ensure valid publications align with expected structure
            valid_publications = [
                pub for pub in publications if len(pub) >= 10
            ]  # Adjust for expected columns
            
            if not valid_publications:
                logger.warning(f"No valid publications found for {canonical_author}. Check data structure.")
            else:
                logger.debug(f"Added {len(valid_publications)} valid publication(s) for {canonical_author}.")
                add_table_to_doc(
                    document,
                    valid_publications,
                    headers=["PubMed ID", "Author", "Year", "Journal", "JID", "Title", "DOI Link", "Citation", "Publication Type", "Access Type"],
                    title="Main Publications",
                )
        else:
            logger.warning(f"No publications found for {canonical_author}.")

        if preprints:
            logger.info(f"> Adding {len(preprints)} preprint(s) for {canonical_author}.")
            
            # Debugging: Check structure of preprints
            logger.debug(f"Raw Preprints for {canonical_author}: {preprints}")
            
            # Adjust for expected structure
            valid_preprints = [
                preprint for preprint in preprints if len(preprint) >= 9
            ]
            
            if not valid_preprints:
                logger.warning(f"No valid preprints found for {canonical_author}. Check data structure.")
            else:
                logger.debug(f"Added {len(valid_preprints)} valid preprint(s) for {canonical_author}.")
                add_table_to_doc(
                    document,
                    valid_preprints,
                    headers=["PubMed ID", "Author", "Year", "Journal", "JID", "Title", "DOI Link", "Citation", "Publication Type"],
                    title="Preprints",
                )
        else:
            logger.warning(f"No preprints found for {canonical_author}.")

        # Summary tables for this author
        add_table_to_doc(
            document,
            [(author, count) for author, count in author_count.items()],
            headers=["Author", "Number of Publications"],
            title="Number of Publications per Author",
        )
        add_table_to_doc(
            document,
            [(year, count) for year, count in year_count.items()],
            headers=["Year", "Number of Publications"],
            title="Number of Publications per Year",
        )
        add_table_to_doc(
            document,
            [(year, journal, count) for year, journals in year_journal_count.items() for journal, count in journals.items()],
            headers=["Year", "Journal", "Number of Publications"],
            title="Number of Publications per Year per Journal",
        )
        add_table_to_doc(
            document,
            [(ptype, year, count) for ptype, years in pub_type_count.items() for year, count in years.items()],
            headers=["Publication Type", "Year", "Number of Publications"],
            title="Publication Type by Year",
        )
        # Add a new page for the next author
        document.add_page_break()
    
    # Add a page break and heading for the combined plotted results
    document.add_page_break()
    document.add_paragraph("Plots of combined results", style='Heading 1')
    # Add combined graphs after individual sections
    logger.info("> Adding graphs for the combined results.")
    for plot_name in [
        f"{output_base_name}_publications_per_author.png",
        f"{output_base_name}_total_publications_preprints_by_author.png",
        f"{output_base_name}_publications_per_year_with_moving_avg.png",
        f"{output_base_name}_publications_per_author_and_year.png",
        f"{output_base_name}_top10_journals_grouped.png",
        f"{output_base_name}_publications_by_access_type.png",
    ]:
        plot_path = os.path.join(results_dir, plot_name)
        if os.path.exists(plot_path):
            try:
                document.add_picture(plot_path, width=Inches(6))
                document.add_paragraph()
            except Exception as e:
                logger.error(f"Error adding plot {plot_path}: {e}")
        else:
            logger.warning(f"Plot file not found: {plot_path}. The actual graph-file in .png-format is not found. Double back and check. For instance, there could to little data to plot.")
            document.add_paragraph(f"Plot {plot_name} not found. The actual graph-file in .png-format is not found. Double back and check. For instance, there could to little data to plot.", style="List Bullet")
    
    # # Add a page break and heading for the collaboration visualizations
    # document.add_page_break()
    # document.add_paragraph("Collaboration Visualizations", style="Heading 1")
    # # Add collaboration visualizations
    # logger.info("> Adding collaboration visualizations.")
    # collab_plot_files = [
    #     f"{output_base_name}_author_collaboration_network.png",
    #     f"{output_base_name}_author_collaboration_heatmap.png",
    #     f"{output_base_name}_top_collaborations_barplot.png",
    #     f"{output_base_name}_author_collaboration_chord.png",
    #     f"{output_base_name}_author_collaboration_matrix.png",
    # ]
    # for plot_file in collab_plot_files:
    #     plot_path = os.path.join(results_dir, plot_file)
    #     if os.path.exists(plot_path):
    #         try:
    #             document.add_picture(plot_path, width=Inches(6))
    #             document.add_paragraph()
    #         except Exception as e:
    #             logger.error(f"Error adding plot {plot_file}: {e}")
    #     else:
    #         logger.warning(f"Collaborative plot file not found: {plot_file}. The actual graph-file in .png-format is not found. Double back and check. For instance, there could to little data to plot.")
    #         document.add_paragraph(f"Collaborative plot {plot_file} not found. The actual graph-file in .png-format is not found. Double back and check. For instance, there could to little data to plot.", style="List Bullet")
    # # Add group collaboration visualizations
    # for plot_name in [
    #     f"{output_base_name}_group_size_distribution.png",
    #     f"{output_base_name}_group_size_by_year.png",
    #     f"{output_base_name}_author_group_heatmap.png",
    # ]:
    #     plot_path = f"{results_dir}/{plot_name}"
    #     if os.path.exists(plot_path):
    #         document.add_picture(plot_path, width=Inches(6))
    #         document.add_paragraph()
    #     else:
    #         logger.warning(f"Collaborative group plot {plot_name} not found. The actual graph-file in .png-format is not found. Double back and check. For instance, there could to little data to plot.")
    #         document.add_paragraph(f"Collaborative group plot {plot_file} not found. The actual graph-file in .png-format is not found. Double back and check. For instance, there could to little data to plot.", style="List Bullet")
    
    # Save the document
    output_path = os.path.join(results_dir, f"{output_base_name}.docx")
    document.save(output_path)
    logger.info(f"Word document saved to [{output_path}].")

####################################################################################################
#                                   PLOT RESULTS                                                   #
####################################################################################################

# Plot the results
def plot_results(author_data, results_dir, logger, output_base_name):
    """
    Plot the results for each author and save the plots.
    """
    date_str = datetime.now().strftime('%Y%m%d')
    plot_filenames = {
        "publications_per_author": f"{output_base_name}_publications_per_author.png",
        "total_publications_preprints_by_author": f"{output_base_name}_total_publications_preprints_by_author.png",
        "publications_per_year_with_moving_avg": f"{output_base_name}_publications_per_year_with_moving_avg.png",
        "publications_per_author_and_year": f"{output_base_name}_publications_per_author_and_year.png",
        "top10_journals_grouped": f"{output_base_name}_top10_journals_grouped.png",
        "publications_by_access_type": f"{output_base_name}_publications_by_access_type.png",
    }

    # Consistent color mapping for authors
    canonical_authors = list(author_data.keys())
    colors = plt.colormaps["tab10"](np.linspace(0, 1, len(canonical_authors)))
    color_map = {canonical_author: colors[idx] for idx, canonical_author in enumerate(canonical_authors)}

    # Access type color mapping
    access_color_map = {"open access": "green", "closed access": "red"}

    # PLOT publications per author
    logger.info("> Plotting publications per author.")
    fig, ax = plt.subplots()
    for canonical_author, (_, _, author_count, _, _, _, _) in author_data.items():
        ax.bar(canonical_author, author_count[canonical_author], color=color_map[canonical_author], label=canonical_author)
    ax.set_xlabel("Authors")
    ax.set_ylabel("Number of Publications")
    ax.set_title("Publications per Author")
    ax.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(os.path.join(results_dir, plot_filenames["publications_per_author"]))

    # PLOT Total number of publications and preprints grouped by author and year (two panels)
    logger.info("> Plotting total publications and preprints grouped by author and year (two panels).")
    fig, axes = plt.subplots(1, 2, figsize=(16, 7), sharey=True)
    access_types = ["open access", "closed access"]

    # Gather all unique years across both panels
    all_years = sorted(set(
        int(pub[2])  # Extract the year
        for _, (publications, preprints, _, _, _, _, _) in author_data.items()
        for pub in publications + preprints
    ))

    # Width of bars for each author in a grouped bar chart
    bar_width = 0.8 / len(author_data)  # Divide the total width by the number of authors

    # Process data for plotting
    for idx, access_type in enumerate(access_types):
        ax = axes[idx]
        for author_idx, (canonical_author, (publications, preprints, _, _, _, _, _)) in enumerate(author_data.items()):
            yearly_totals = defaultdict(int)
            for pub in publications + preprints:
                pub_access_type = pub[-1]  # Access type is the last field
                year = int(pub[2])  # Year is the third field
                if pub_access_type == access_type:
                    yearly_totals[year] += 1
            
            # Compute bar positions for the author within each year
            positions = [x + author_idx * bar_width for x in range(len(all_years))]
            counts = [yearly_totals.get(year, 0) for year in all_years]  # Ensure all years are represented
            
            # Plot the bars for the current author
            ax.bar(
                positions,
                counts,
                bar_width,
                label=canonical_author,
                color=color_map[canonical_author],
                alpha=0.8,
            )
        
        # Configure the x-axis
        ax.set_title(f"Total Publications ({access_type.capitalize()})")
        ax.set_xlabel("Year")
        ax.set_xticks([x + (bar_width * len(author_data)) / 2 for x in range(len(all_years))])  # Center ticks for years
        ax.set_xticklabels(all_years, rotation=45)  # Use unified years across panels
        
        if idx == 0:
            ax.set_ylabel("Total Publications and Preprints")
        ax.legend(bbox_to_anchor=(1.05, 1), loc="upper left")

    plt.tight_layout()
    plt.savefig(os.path.join(results_dir, plot_filenames["total_publications_preprints_by_author"]))

    # PLOT publications per year colored and grouped by author with moving average
    logger.info("> Plotting publications per year with moving average.")
    fig, ax = plt.subplots(figsize=(10, 6))
    width = 0.8 / len(author_data)  # Divide bar width by number of authors for grouped bars

    # Gather all unique years across authors
    all_years = sorted(set(year for _, (_, _, _, year_count, _, _, _) in author_data.items() for year in year_count))

    # Define the moving average window size
    moving_avg_window = 3

    # Plot each author's publications per year with bars and moving average
    for idx, (canonical_author, (_, _, _, year_count, _, _, _)) in enumerate(author_data.items()):
        counts = [year_count.get(year, 0) for year in all_years]  # Publications count for each year
        bar_positions = [year + idx * width for year in range(len(all_years))]  # Bar positions for this author

        # Plot bars for this author
        ax.bar(
            bar_positions,
            counts,
            width=width,
            color=color_map[canonical_author],
            label=canonical_author,
            alpha=0.8,
        )

        # Calculate the moving average (adjust at edges)
        moving_avg = []
        for i in range(len(counts)):
            window_start = max(0, i - moving_avg_window + 1)
            window = counts[window_start:i + 1]
            moving_avg.append(sum(window) / len(window))

        # Plot the moving average line
        ax.plot(
            range(len(moving_avg)),
            moving_avg,
            marker='o',
            label=f"{canonical_author} (Moving Avg)",
            linestyle='--',
            color=color_map[canonical_author],
        )

    # Customize the x-axis
    ax.set_xticks(range(len(all_years)))
    ax.set_xticklabels(all_years, rotation=45)
    ax.set_xlabel("Year")
    ax.set_ylabel("Number of Publications")
    ax.set_title("Publications per Year with Moving Average")
    ax.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.tight_layout()

    # Save the plot
    plt.savefig(os.path.join(results_dir, plot_filenames["publications_per_year_with_moving_avg"]))

    # PLOT publications per author per year (stacked bar plot)
    logger.info("> Plotting publications per author and year.")
    fig, ax = plt.subplots()
    width = 0.2
    x = sorted({year for _, (_, _, _, year_count, _, _, _) in author_data.items() for year in year_count})
    x_indices = np.arange(len(x))
    for idx, (canonical_author, (_, _, _, year_count, author_year_count, _, _)) in enumerate(author_data.items()):
        counts = [author_year_count[canonical_author].get(year, 0) for year in x]
        ax.bar(x_indices + idx * width, counts, width, color=color_map[canonical_author], label=canonical_author)
    ax.set_xticks(x_indices + width / 2 * (len(canonical_authors) - 1))
    ax.set_xticklabels(x, rotation=45)
    ax.set_xlabel("Year")
    ax.set_ylabel("Number of Publications")
    ax.set_title("Publications per Author and Year")
    ax.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.tight_layout()
    plt.savefig(os.path.join(results_dir, plot_filenames["publications_per_author_and_year"]))

    # PLOT Top 10 journals by number of publications (grouped by year)
    logger.info("> Plotting top 10 journals.")
    journal_counts = defaultdict(lambda: defaultdict(int))
    for _, (_, _, _, _, _, year_journal_count, _) in author_data.items():
        for year, journals in year_journal_count.items():
            for journal, count in journals.items():
                journal_counts[journal][year] += count

    # Flatten journal counts and sort by total
    total_journal_counts = {journal: sum(counts.values()) for journal, counts in journal_counts.items()}
    sorted_journals = sorted(total_journal_counts.items(), key=lambda x: x[1], reverse=True)
    top_10_journals = [journal for journal, _ in sorted_journals[:10]]  # Keep only top 10

    # Prepare data for plotting
    all_years = sorted({year for journal in journal_counts for year in journal_counts[journal]})
    grouped_counts = defaultdict(lambda: defaultdict(int))
    for journal in top_10_journals:
        for year in all_years:
            grouped_counts[journal][year] = journal_counts[journal].get(year, 0)

    # Plot data
    fig, ax = plt.subplots(figsize=(12, 7))
    bar_width = 0.8 / len(top_10_journals)  # Adjust bar width based on the number of journals
    colors = plt.cm.tab10(np.linspace(0, 1, len(top_10_journals)))  # Generate a color palette

    # Create x positions for each year
    x_positions = np.arange(len(all_years))

    # Plot each journal as a separate group
    for idx, journal in enumerate(top_10_journals):
        counts = [grouped_counts[journal][year] for year in all_years]
        offset = idx * bar_width  # Offset each journal's bars
        ax.bar(
            x_positions + offset,
            counts,
            width=bar_width,
            label=journal,
            color=colors[idx],
        )

    # Customize x-axis and labels
    ax.set_xticks(x_positions + bar_width * len(top_10_journals) / 2)  # Center x-ticks
    ax.set_xticklabels(all_years, rotation=45)
    ax.set_xlabel("Year")
    ax.set_ylabel("Number of Publications")
    ax.set_title("Top 10 Journals (Grouped by Year)")
    ax.legend(bbox_to_anchor=(1.05, 1), loc="upper left")

    plt.tight_layout()
    plt.savefig(os.path.join(results_dir, plot_filenames["top10_journals_grouped"]))

    # PLOT Publications grouped by access type and year (not stacked)
    logger.info("> Plotting publications by access type and year (grouped).")
    fig, ax = plt.subplots(figsize=(12, 7))
    access_counts = {"open access": defaultdict(int), "closed access": defaultdict(int)}

    # Process data for grouped bar plot
    for _, (publications, _, _, _, _, _, _) in author_data.items():
        for pub in publications:
            access_type = pub[-1]
            if access_type in access_counts:
                access_counts[access_type][int(pub[2])] += 1

    all_years = sorted(set(year for totals in access_counts.values() for year in totals.keys()))
    x_indices = np.arange(len(all_years))
    width = 0.35
    for idx, (access_type, yearly_counts) in enumerate(access_counts.items()):
        counts = [yearly_counts[year] for year in all_years]
        ax.bar(
            x_indices + idx * width,
            counts,
            width=width,
            label=access_type,
            color=access_color_map[access_type],
            alpha=0.8,
        )
    ax.set_xticks(x_indices + width / 2)
    ax.set_xticklabels(all_years, rotation=45)
    ax.set_xlabel("Year")
    ax.set_ylabel("Number of Publications")
    ax.set_title("Publications by Access Type and Year (Grouped)")
    ax.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.tight_layout()
    plt.savefig(os.path.join(results_dir, plot_filenames["publications_by_access_type"]))

# Create a network graph for author collaborations
def create_network_graph(collaboration_data, results_dir, output_file):
    """
    Create a network graph showing collaborations between authors.
    Args:
        collaboration_data (list): List of tuples (author1, author2, num_collaborations).
        output_file (str): File path to save the plot.
    """
    G = nx.Graph()
    
    # Add edges with weights
    for author1, author2, num_papers in collaboration_data:
        G.add_edge(author1, author2, weight=num_papers)
    
    # Draw the graph
    pos = nx.spring_layout(G)  # Positioning of nodes
    weights = [G[u][v]['weight'] for u, v in G.edges()]
    nx.draw(
        G, pos, with_labels=True, width=weights, node_color="skyblue",
        edge_color="gray", node_size=2000, font_size=10
    )
    plt.title("Author Collaboration Network")
    plt.savefig(output_file)
    plt.close()

# Create a heatmap for collaborations
def create_heatmap(collaboration_matrix, authors, results_dir, output_file):
    """
    Create a heatmap showing collaborations between authors.
    Args:
        collaboration_matrix (np.array): Collaboration matrix (2D array).
        authors (list): List of authors corresponding to rows/columns.
        output_file (str): File path to save the plot.
    """
    df = pd.DataFrame(collaboration_matrix, index=authors, columns=authors)
    sns.heatmap(df, annot=True, fmt="d", cmap="Blues", cbar=True)
    plt.title("Author Collaboration Heatmap")
    plt.xlabel("Authors")
    plt.ylabel("Authors")
    plt.savefig(output_file)
    plt.close()

# Create a bar plot of collaborations
def create_bar_plot(collaboration_data, results_dir, output_path):
    """
    Create a bar plot for the top collaborations.

    Args:
        collaboration_data (list): List of tuples (author1, author2, num_collaborations).
        output_path (str): Path to save the plot.

    Returns:
        None
    """
    # Convert collaboration_data into a DataFrame
    collaboration_df = pd.DataFrame(collaboration_data, columns=["Author1", "Author2", "NumCollaborations"])

    # Combine author pairs for display purposes
    collaboration_df["Pair"] = collaboration_df.apply(
        lambda row: f"{row['Author1']} & {row['Author2']}", axis=1
    )

    # Sort by number of collaborations and select the top 10
    top_collaborations = collaboration_df.sort_values(by="NumCollaborations", ascending=False).head(10)

    # Plot
    plt.figure(figsize=(10, 6))
    plt.barh(top_collaborations["Pair"], top_collaborations["NumCollaborations"], color="skyblue")
    plt.xlabel("Number of Collaborations")
    plt.ylabel("Author Pairs")
    plt.title("Top Collaborations")
    plt.gca().invert_yaxis()  # Invert y-axis for better readability
    plt.tight_layout()

    # Save the plot
    plt.savefig(output_path)
    plt.close()

# Create a chord diagram for collaborations
hv.extension("bokeh")
def create_chord_diagram(collaboration_data, results_dir, output_base_name, logger, save_html=True, save_png=False):
    """
    Create a chord diagram for author collaborations.

    Args:
        collaboration_data (list): List of tuples (author1, author2, num_collaborations).
        results_dir (str): Directory to save results.
        output_base_name (str): Base name for output files.
        logger: Logger instance.
        save_html (bool): Whether to save as an interactive HTML file.
        save_png (bool): Whether to save as a static PNG image.
    """
    from holoviews.plotting.util import dim
    import holoviews as hv
    from holoviews import opts
    hv.extension("bokeh")

    # Validate collaboration data
    if not collaboration_data or not all(len(row) == 3 for row in collaboration_data):
        logger.error("Invalid collaboration data. Ensure it's a list of (author1, author2, num_collaborations) tuples.")
        return

    try:
        logger.info("Generating Chord Diagram.")

        # Create DataFrame for validation
        df = pd.DataFrame(collaboration_data, columns=["Author 1", "Author 2", "Collaborations"])
        logger.debug(f"Collaboration data for chord diagram:\n{df.head()}")

        # Create Chord diagram
        chord = hv.Chord((df, hv.Dataset(df, ["Author 1", "Author 2"], "Collaborations")))
        chord.opts(
            opts.Chord(cmap="Category20", edge_color=dim("Collaborations"), node_size=10, labels="index", width=800, height=800)
        )

        # Save as HTML
        if save_html:
            html_path = os.path.join(results_dir, f"{output_base_name}_author_collaboration_chord.html")
            hv.save(chord, html_path, fmt="html")
            logger.info(f"Chord Diagram saved as HTML to {html_path}.")

        # Save as PNG
        if save_png:
            png_path = os.path.join(results_dir, f"{output_base_name}_author_collaboration_chord.png")
            import panel as pn
            pn.extension()  # Required for PNG rendering
            hv.save(chord, png_path, fmt="png")
            logger.info(f"Chord Diagram saved as PNG to {png_path}.")
    except Exception as e:
        logger.error(f"Failed to generate Chord Diagram: {e}")

# Create a matrix diagram for collaborations
def create_matrix_diagram(collaboration_matrix, authors, results_dir, output_file):
    """
    Create a matrix diagram showing collaborations between authors.
    Args:
        collaboration_matrix (np.array): Collaboration matrix (2D array).
        authors (list): List of authors corresponding to rows/columns.
        output_file (str): File path to save the plot.
    """
    fig, ax = plt.subplots()
    cax = ax.matshow(collaboration_matrix, cmap="Blues")
    fig.colorbar(cax)
    ax.set_xticks(np.arange(len(authors)))
    ax.set_yticks(np.arange(len(authors)))
    ax.set_xticklabels(authors, rotation=90)
    ax.set_yticklabels(authors)
    plt.title("Collaboration Matrix")
    plt.savefig(output_file)
    plt.close()

# Create visualizations for group collaborations
def create_collaboration_visualizations(group_collaboration_data, authors, results_dir, output_base_name, logger):
    if not group_collaboration_data:
        logger.warning("Insufficient data for collaboration visualizations.")
        return
    
    logger.info("Creating collaboration visualizations.")
    try:
        plot_group_size_distribution(group_collaboration_data, f"{results_dir}/{output_base_name}_group_size_distribution.png")
        logger.info("> Group Size Distribution plot saved.")
        
        plot_group_chord_diagram(group_collaboration_data, authors, f"{results_dir}/{output_base_name}_group_chord_diagram.html")
        logger.info("> Group Chord Diagram saved.")
        
        plot_group_size_by_year(group_collaboration_data, f"{results_dir}/{output_base_name}_group_size_by_year.png")
        logger.info("> Group Size by Year plot saved.")
        
        plot_author_group_heatmap(group_collaboration_data, authors, f"{results_dir}/{output_base_name}_author_group_heatmap.png")
        logger.info("> Author Group Heatmap plot saved.")
        
        plot_interactive_author_network(group_collaboration_data, authors, f"{results_dir}/{output_base_name}_author_network.html")
        logger.info("> Interactive Author Network plot saved.")
    except Exception as e:
        logger.error(f"Error while creating collaboration visualizations: {e}")

# Group collabborations: Bar Plot, Group Sizes vs. Frequency
def plot_group_size_distribution(group_collaboration_data, output_file):
    group_sizes = [len(group) for group in group_collaboration_data.keys()]
    frequencies = list(group_collaboration_data.values())
    
    plt.figure(figsize=(8, 6))
    plt.bar(group_sizes, frequencies, color='skyblue')
    plt.xlabel("Group Size")
    plt.ylabel("Number of Collaborations")
    plt.title("Collaboration Group Size Distribution")
    plt.xticks(range(2, max(group_sizes) + 1))  # Only show valid group sizes
    plt.tight_layout()
    plt.savefig(output_file)
    plt.close()

# Group collabborations: Chord Diagram
def plot_group_chord_diagram(group_collaboration_data, author_index, output_file):
    chord_data = []
    for group, count in group_collaboration_data.items():
        for i, author1 in enumerate(group):
            for author2 in group[i + 1:]:
                chord_data.append((author1, author2, count))
    
    chord_diagram = hv.Chord(chord_data).opts(
        opts.Chord(labels='index', cmap='Category20', edge_cmap='viridis',
                   edge_color=dim('value'), node_size=10, height=600, width=600)
    )
    hv.save(chord_diagram, output_file, fmt='html')

# Group collabborations: Stacked Bar Chart, Group Size vs. Year
def plot_group_size_by_year(group_collaboration_data, output_file):
    rows = []
    for group, count in group_collaboration_data.items():
        group_size = len(group)
        year = group[-1]['year']  # Assuming year metadata is in the group
        rows.append({'Year': year, 'Group Size': group_size, 'Count': count})
    df = pd.DataFrame(rows)

    pivot_df = df.pivot_table(index='Year', columns='Group Size', values='Count', aggfunc='sum', fill_value=0)

    pivot_df.plot(kind='bar', stacked=True, figsize=(10, 6), cmap='viridis')
    plt.xlabel("Year")
    plt.ylabel("Number of Collaborations")
    plt.title("Group Size by Year")
    plt.tight_layout()
    plt.savefig(output_file)
    plt.close()

# Group collabborations: Heatmap, Group Members vs. Group Size
def plot_author_group_heatmap(group_collaboration_data, authors, output_file):
    author_group_counts = {author: {size: 0 for size in range(2, 6)} for author in authors}
    for group, count in group_collaboration_data.items():
        size = len(group)
        for author in group:
            author_group_counts[author][size] += count

    heatmap_df = pd.DataFrame(author_group_counts).T
    heatmap_df = heatmap_df.fillna(0)

    plt.figure(figsize=(12, 8))
    sns.heatmap(heatmap_df, cmap="YlGnBu", annot=True, fmt="d")
    plt.xlabel("Group Size")
    plt.ylabel("Author")
    plt.title("Author Participation by Group Size")
    plt.tight_layout()
    plt.savefig(output_file)
    plt.close()

# Group collabborations: Interactive Network Graph
def plot_interactive_author_network(group_collaboration_data, authors, output_file):
    G = nx.Graph()
    for group, count in group_collaboration_data.items():
        for i, author1 in enumerate(group):
            for author2 in group[i + 1:]:
                if G.has_edge(author1, author2):
                    G[author1][author2]['weight'] += count
                else:
                    G.add_edge(author1, author2, weight=count)
    
    pos = nx.spring_layout(G)
    edge_x, edge_y, edge_width = [], [], []
    for edge in G.edges(data=True):
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        edge_x.append(x0)
        edge_x.append(x1)
        edge_x.append(None)
        edge_y.append(y0)
        edge_y.append(y1)
        edge_y.append(None)
        edge_width.append(edge[2]['weight'])

    edge_trace = go.Scatter(x=edge_x, y=edge_y, line=dict(width=edge_width, color='Gray'), hoverinfo='none', mode='lines')
    node_trace = go.Scatter(x=[pos[n][0] for n in G.nodes], y=[pos[n][1] for n in G.nodes],
                             mode='markers+text', text=list(G.nodes),
                             marker=dict(size=10, color='LightBlue'))

    fig = go.Figure(data=[edge_trace, node_trace])
    fig.write_html(output_file)

####################################################################################################
#                                   MAIN FUNCTION                                                  #
####################################################################################################

# Main function
def main():
    args = parse_arguments()

    # Get today's date
    today = datetime.now().strftime('%Y%m%d')

    # Make sure the results directory exists
    results_dir = "results"
    os.makedirs(results_dir, exist_ok=True)

    # File base naming convention
    base_name = args.output_file if args.output_file else "CDL_UMCU_Publications"
    output_base_name = f"{today}_{base_name}"
    
    # Set up logging
    logger = setup_logger(results_dir, output_base_name, args.verbose, args.debug)

    # Ensure all required packages are installed
    for package in ['Bio', 'docx', 'matplotlib', 'numpy', 'pandas', 'seaborn', 'networkx', 'holoviews']:
        check_install_package(package, logger)

    # Set year range if provided
    start_year, end_year = None, None
    if args.year:
        start_year, end_year = parse_year_range(args.year)

    # Check if email or year is provided when using dummy data
    if args.dummy and (args.email or args.year):
        logger.warning("The --dummy flag is set; ignoring --email and --year arguments.")

    # Print some information
    logger.info(f"Running {VERSION_NAME} v{VERSION} ({VERSION_DATE})\n")

    # # Generate or fetch data
    if args.dummy:
        logger.info("*** Using dummy data for testing. ***\n")
        author_data = generate_dummy_data()
        logger.info(f"Generated dummy data for {len(author_data)} authors.")
        for author, (publications, preprints, *_) in author_data.items():
            logger.info(f"> {author}: {len(publications)} publications, {len(preprints)} preprints.")
    else:
        # Set the email for Entrez
        Entrez.email = args.email

        # Print some information
        logger.info(f"Settings:")
        logger.info(f"> Search parameters given:")
        logger.info(f"  - authors: {args.names}")
        logger.info(f"  - department(s): {args.departments}" if not args.ignore_departments else "  - no department filter used.")
        logger.info(f"  - organization: {args.organization}")
        logger.info(f"  - filtering by year (range) [{args.year}]" if args.year else "  - no year filter used.")
        logger.info(f"  - output file(s): [{output_base_name}]")
        logger.info(f"> PubMed email used: {args.email}. Note that this is only used by PubMed to log the queries, it is not used for the actual query.")
        logger.info(f"> Debug mode: {'On' if args.debug else 'Off'}.")
        logger.info(f"> Verbose mode: {'On' if args.verbose else 'Off'}.\n")
        

        # Collect all PubMed IDs for the given author(s)
        logger.info(f"Querying PubMed for publications and preprints.\n")
        author_data = {}
        for main_author in args.names:
            all_pubmed_ids = set()  # Use a set to collect unique IDs
            # Retrieve the canonical author and their aliases
            canonical_author = get_canonical_author(main_author, logger)
            canonical_author_aliases = ALIAS_MAPPING.get(canonical_author, [canonical_author])

            # Construct the author query based on whether it's an ORCID
            author_query = " OR ".join(
                f'({alias}[Author - Identifier])' if re.match(r"^\d{4}-\d{4}-\d{4}-\d{4}$", alias)
                else f'({alias}[Author])'
                for alias in canonical_author_aliases
            )

            logger.debug(f"Searching PubMed for canonical author '{canonical_author}' with aliases: {canonical_author_aliases}.")

            # Department queries
            department_query = ""
            if not args.ignore_departments:
                department_queries = []
                for department in args.departments:
                    department_aliases = DEPARTMENT_ALIAS_MAPPING.get(department, [department])
                    department_queries.append(" OR ".join(f'({alias}[Affiliation])' for alias in department_aliases))
                department_query = " OR ".join(f"({query})" for query in department_queries)

                logger.debug(f"> Combined department query: {department_query}")

            # Build organization queries
            organization_queries = []
            for organization in args.organization:
                organization_aliases = ORGANIZATION_ALIAS_MAPPING.get(organization, [organization])
                organization_queries.append(" OR ".join(f'({alias}[Affiliation])' for alias in organization_aliases))

            # Combine organization queries
            organization_query = " OR ".join(f"({query})" for query in organization_queries)
            logger.debug(f"> Combined organization query: {organization_query}")


            # Construct the full search query
            if args.ignore_departments:
                search_query = f"(({author_query}) AND ({organization_query}))"
            else:
                search_query = f"(({author_query}) AND ({department_query})) AND ({organization_query})"
            logger.debug(f"Constructed PubMed search query: {search_query}")

            try:
                record = fetch_with_retry(db="pubmed", term=search_query)
                if not record:
                    logger.error(f"No data returned for query [{search_query}]. Skipping.")
                    continue

                all_pubmed_ids.update(record["IdList"])  # Add unique PubMed IDs to the set
            # Create exception when no publications are found for the author
            except Exception as e:
                logger.error(f"Error querying PubMed for author '{canonical_author}' [search query: {search_query}]: {e}")
                logger.info(f"Skipping author '{canonical_author}'.")
                continue  # Ensure the loop continues for the next author
            
            # Validate the results to ensure all criteria are met 
            # See below the whole script for the function validate_results

            # Log the number of unique IDs for this author
            logger.info(f"Found {len(all_pubmed_ids)} unique publications for author '{canonical_author}'.")

            # Fetch detailed publication data for the current author
            publications, preprints = fetch_publication_details(
                sorted(all_pubmed_ids), logger, canonical_author, start_year, end_year
            )
            author_count, year_count, author_year_count, year_journal_count, pub_type_count = analyze_publications(
                publications, canonical_author, logger
            )
            author_data[canonical_author] = (publications, preprints, author_count, year_count, author_year_count, year_journal_count, pub_type_count)

            logger.info(f"Parsed {len(publications)} publications and {len(preprints)} preprints for [{canonical_author}].\n")

    # Summarizing and saving results
    logger.info(f"Done. Summarizing and saving results.\n")
    logger.info(f"Saving plots to [{results_dir}].")
    plot_results(author_data, results_dir, logger, output_base_name)

    # Process collaboration data
    logger.info(f"Processing collaboration data.")
    logger.debug(f"Author data: {list(author_data.items())[:5]}")  # Debug first 5 rows from author_data

    # Get collaboration data, group collaboration data, and details
    collaboration_data, group_collaboration_data, collaboration_matrix, authors = process_collaborations(author_data, logger, ALIAS_MAPPING, results_dir, output_base_name, Entrez.email, max_group_size=5)
    
    # Check validity of collaboration data
    if not collaboration_data or not all(len(row) == 3 for row in collaboration_data):
        logger.error("Invalid collaboration data. Ensure it's a list of (author1, author2, num_collaborations) tuples.")
    return

    # Create collaboration visualizations if data is available
    if not collaboration_data:
        logger.warning("No collaboration data available; skipping collaboration visualizations.")
    else:
        logger.info(f"Found {len(collaboration_data)} unique collaborations.")
        # count unique collaborations
        logger.info("Found {} unique collaborations.".format(len(collaboration_data)))
        logger.debug(f"Collaboration data: {collaboration_data[:5]}")  # Debug first 5 collaborations
        logger.debug(f"Authors: {authors}")
        if args.debug:
            # save collaboration data as python objects
            logger.debug(f"Saving collaboration data for debug purposes.")
            collaboration_data_path = os.path.join(results_dir, f"{output_base_name}_collaboration_data.pkl")
            group_collaboration_data_path = os.path.join(results_dir, f"{output_base_name}_group_collaboration_data.pkl")
            collaboration_matrix_path = os.path.join(results_dir, f"{output_base_name}_collaboration_matrix.pkl")
            authors_path = os.path.join(results_dir, f"{output_base_name}_authors.pkl")
            with open(collaboration_data_path, "wb") as f:
                pickle.dump(collaboration_data, f)
            with open(group_collaboration_data_path, "wb") as f:
                pickle.dump(group_collaboration_data, f)
            with open(collaboration_matrix_path, "wb") as f:
                pickle.dump(collaboration_matrix, f)
            with open(authors_path, "wb") as f:
                pickle.dump(authors, f)

        # Generate visualizations for pairwise collaborations
        logger.info(f"Generating pairwise collaboration visualizations.\n")
        create_network_graph(collaboration_data, results_dir, f"{output_base_name}_author_collaboration_network.png")
        create_bar_plot(collaboration_data, results_dir, f"{output_base_name}_top_collaborations_barplot.png")
        create_chord_diagram(collaboration_data, results_dir, output_base_name, logger, save_html=True, save_png=True)
    
        # Generate collaboration matrix visualizations
        if collaboration_matrix is None or collaboration_matrix.size == 0:
            logger.warning("Collaboration matrix is empty; skipping visualizations.")
        else: 
            logger.info(f"Collaboration matrix was created. Generating associated visualizations.")
            logger.debug(f"Collaboration matrix shape: {collaboration_matrix.shape}")
            create_heatmap(collaboration_matrix, authors, results_dir, f"{output_base_name}_author_collaboration_heatmap.png")
            create_matrix_diagram(collaboration_matrix, authors, results_dir, f"{output_base_name}_author_collaboration_matrix.png")
    
    # Generate group collaboration visualizations
    if group_collaboration_data is None or len(group_collaboration_data) == 0:
        logger.warning("Group collaboration data is empty; skipping group collaboration visualizations.")
    else:
        logger.info(f"Group collaboration data was created. Generating visualizations.")
        create_collaboration_visualizations(group_collaboration_data, authors, results_dir, output_base_name, logger)
        
    # Save results to Word and Excel
    write_to_word(author_data, output_base_name, results_dir, logger, args)
    write_to_excel(author_data, output_base_name, results_dir, logger)

    logger.info(f"Saved the following results:")
    logger.info(f"> Data summarized and saved to {os.path.join(results_dir, f'{output_base_name}.docx')}.")
    logger.info(f"> Excel concatenated and saved to {os.path.join(results_dir, f'{output_base_name}.xlsx')}.")
    
    logger.info(f"> Plots saved to {results_dir}/.")
    logger.info(f"> Log file saved to {os.path.join(results_dir, f'{output_base_name}.log')}.\n")
    logger.info(f"Thank you for using {VERSION_NAME} v{VERSION} ({VERSION_DATE}).")
    logger.info(f"{COPYRIGHT}")
    logger.info(f"Script completed successfully on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.")

if __name__ == "__main__":
    main()
